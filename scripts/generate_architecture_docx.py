from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_page_margins(section, margin_inch=0.6):
    section.top_margin = Inches(margin_inch)
    section.bottom_margin = Inches(margin_inch)
    section.left_margin = Inches(margin_inch)
    section.right_margin = Inches(margin_inch)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, title, items, title_color="10223A"):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(title_color)

    for item in items:
        paragraph = cell.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        run.font.size = Pt(10.5)


def add_heading(document, text, size, bold=True, color="10223A", align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def add_body(document, text, size=10.5, color="425466", align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def generate_docx(output_path: Path):
    document = Document()
    section = document.sections[0]
    set_page_margins(section)

    add_heading(document, "업비트 자동매매 시스템 아키텍처", 22)
    add_body(
        document,
        "로컬은 연구와 관제, VPS는 실거래 실행을 담당하는 구조로 정리한 시스템 개요 문서.",
        size=11,
    )

    table = document.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    table.autofit = True

    local_items = [
        "PM Orchestrator와만 소통",
        "전략 연구 및 백테스트",
        "코드 수정, Git 관리, 배포",
        "대시보드와 로그 확인",
    ]
    vps_items = [
        "종목 스캐너와 신호 엔진",
        "주문 엔진과 포지션 관리자",
        "리스크 가드와 손실 제한",
        "WebSocket 리스너와 스케줄러",
        "API Key / 환경변수 보관",
    ]
    upbit_items = [
        "Quotation API",
        "Exchange API",
        "실시간 WebSocket",
    ]
    runtime_items = [
        "장애 감지",
        "하트비트 체크",
        "실행 로그 저장",
        "일일 손익 집계",
    ]
    monitor_items = [
        "Telegram 즉시 알림",
        "체결 / 장애 / 하트비트 통지",
        "사용자 상태 보고",
    ]
    rules_items = [
        "1회 최대손실 1만원",
        "1일 최대손실 3만원",
        "동시 보유 종목 3개",
        "실거래 키는 VPS에서만 사용",
    ]

    set_cell_text(table.cell(0, 0), "사용자 / 로컬 노트북", local_items)
    shade_cell(table.cell(0, 0), "E4F2FF")
    set_cell_text(table.cell(0, 1), "VPS 실거래 서버", vps_items)
    shade_cell(table.cell(0, 1), "DFF5E7")
    set_cell_text(table.cell(0, 2), "업비트", upbit_items)
    shade_cell(table.cell(0, 2), "FFE9DC")

    set_cell_text(table.cell(1, 0), "운영 런타임", runtime_items)
    shade_cell(table.cell(1, 0), "FFF5DC")
    set_cell_text(table.cell(1, 1), "모니터링 / 알림", monitor_items)
    shade_cell(table.cell(1, 1), "FFE2E7")
    set_cell_text(table.cell(1, 2), "리스크 기본값", rules_items)
    shade_cell(table.cell(1, 2), "F1F4F8")

    document.add_paragraph("")
    add_heading(document, "주요 흐름", 15)
    for flow in [
        "배포 / 제어: 로컬 노트북에서 VPS로 코드를 배포하고 운영 상태를 관리한다.",
        "시세 / 주문: VPS가 업비트 시세를 수신하고 조건 충족 시 주문을 실행한다.",
        "알림 / 보고: 체결, 장애, 손익 상태를 Telegram 등으로 사용자에게 전달한다.",
    ]:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(flow)
        run.font.size = Pt(10.5)

    add_heading(document, "운영 원칙", 15)
    for principle in [
        "로컬은 연구·백테스트·관제 중심으로 사용한다.",
        "VPS는 고정 공인 IPv4 기반으로 실주문만 담당한다.",
        "업비트 API Key는 VPS 환경에만 저장한다.",
        "노트북 단독 실거래 운영은 유동 IP 문제로 권장하지 않는다.",
    ]:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(principle)
        run.font.size = Pt(10.5)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    generate_docx(Path("workspace/reports/bitcoin-trading-system-architecture.docx"))
